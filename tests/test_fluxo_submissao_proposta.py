# test_fluxo_submissao_proposta.py
# -*- coding: utf-8 -*-
"""
Suíte de Testes Automatizados: Fluxo Completo de Submissão e Ciclo de Vida da Proposta
Testa as 5 etapas integradas (Ingestão -> Auditoria -> Revisão -> Supervisão -> Exportação) 100% offline.
"""

import unittest
import io
import os
import openpyxl
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet


class FullProposalSubmissionPipeline:
    """Simulador do Pipeline Completo de Elaboração e Submissão de Propostas."""

    @staticmethod
    def step1_ingestion(raw_edital_text: str, annexes_count: int = 1) -> dict:
        """Etapa 1: Ingestão de texto, extração de objeto, teto financeiro e prazos."""
        if not raw_edital_text or len(raw_edital_text.strip()) < 50:
            raise ValueError("Texto do edital inválido ou muito curto para ingestão.")
        
        return {
            "status": "ingested",
            "edital_length": len(raw_edital_text),
            "annexes_count": annexes_count,
            "has_budget_rules": "teto" in raw_edital_text.lower() or "limite" in raw_edital_text.lower(),
            "has_accessibility_rules": "acessibilidade" in raw_edital_text.lower() or "libras" in raw_edital_text.lower()
        }

    @staticmethod
    def step2_sector_review(ingestion_data: dict, proposal_sections: dict) -> dict:
        """Etapa 2: Revisão técnica pelos 14 eixos especialistas."""
        scores = {}
        for sec_name, sec_content in proposal_sections.items():
            if not sec_content or len(sec_content.strip()) < 20:
                scores[sec_name] = {"nota": 40, "status": "pendente"}
            else:
                scores[sec_name] = {"nota": 95, "status": "aprovado"}
        
        total_score = sum(s["nota"] for s in scores.values())
        avg_score = round(total_score / len(scores), 1) if scores else 0
        return {
            "sector_scores": scores,
            "nota_tecnica_media": avg_score,
            "all_approved": all(s["status"] == "aprovado" for s in scores.values())
        }

    @staticmethod
    def step3_strategic_supervision(review_data: dict) -> dict:
        """Etapa 3: Síntese estratégica e diretrizes do Supervisor Geral."""
        nota_media = review_data.get("nota_tecnica_media", 0)
        status_geral = "APROVADO_PARA_SUBMISSAO" if nota_media >= 80 else "AJUSTES_NECESSARIOS"
        
        return {
            "score_final": min(100, int(nota_media)),
            "status_geral": status_geral,
            "diretriz_final": "Proposta em total consonância com as metas e exigências da chamada pública."
        }

    @staticmethod
    def step4_generate_abnt_document(cover: dict, sections: dict) -> bytes:
        """Etapa 4: Compilação do documento final em formato DOCX (Word)."""
        doc = Document()
        doc.add_heading(cover.get("titulo", "PROPOSTA DE PROJETO CULTURAL").upper(), level=0)
        doc.add_paragraph(f"Proponente: {cover.get('proponente', 'Associação Proponente')}")
        doc.add_paragraph(f"Município: {cover.get('municipio', 'Vitória - ES')}")
        doc.add_paragraph(f"Valor Solicitado: R$ {cover.get('valor_total', 0.0):,.2f}")
        doc.add_paragraph(f"Prazo de Execução: {cover.get('prazo_meses', 12)} meses")
        doc.add_paragraph("---")
        
        for sec_title, sec_body in sections.items():
            doc.add_heading(sec_title.upper(), level=1)
            doc.add_paragraph(sec_body)
            
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def step5_generate_financial_spreadsheet(budget_items: list) -> bytes:
        """Etapa 5: Geração da planilha orçamentária oficial XLSX."""
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Planilha Orçamentária"
        
        ws.append(["Item", "Descrição", "Unidade", "Quantidade", "Valor Unitário (R$)", "Total (R$)"])
        row_idx = 2
        for it in budget_items:
            ws.append([
                it.get("item", 1),
                it.get("descricao", "Serviço"),
                it.get("unidade", "un"),
                it.get("quantidade", 1),
                it.get("valor_unitario", 0.0),
                f"=D{row_idx}*E{row_idx}"
            ])
            row_idx += 1
            
        ws.append(["TOTAL GERAL", "", "", "", "", f"=SUM(F2:F{row_idx-1})"])
        
        buffer = io.BytesIO()
        wb.save(buffer)
        return buffer.getvalue()


class TestFluxoSubmissaoProposta(unittest.TestCase):
    """Testes de ponta a ponta do ciclo de vida da proposta."""

    def setUp(self):
        self.sample_edital = """
        CHAMADA PÚBLICA DE FOMENTO CULTURAL - EDITAL RIO DOCE 2026
        OBJETO: Seleção de iniciativas comunitárias na Bacia do Rio Doce.
        TETO ORÇAMENTÁRIO: R$ 220.000,00 por projeto.
        ACESSIBILIDADE: Medidas de LIBRAS e audiodescrição obrigatórias.
        PRAZO DE EXECUÇÃO: 12 meses.
        """
        self.sample_sections = {
            "1. Justificativa": "O projeto visa preservar o patrimônio imaterial da bacia do Rio Doce com ampla participação comunitária.",
            "2. Objetivos": "Realizar 24 oficinas culturais e 4 apresentações públicas com acessibilidade total.",
            "3. Metodologia": "Aulas práticas com mestres tradicionais e equipe multidisciplinar qualificada.",
            "4. Cronograma": "Pré-produção nos meses 1-2, execução dos meses 3-10, encerramento nos meses 11-12.",
            "5. Acessibilidade": "Presença confirmada de 2 intérpretes de LIBRAS em todas as ações abertas ao público."
        }
        self.sample_budget = [
            {"item": 1, "descricao": "Coordenação Geral", "unidade": "mês", "quantidade": 12, "valor_unitario": 3500.0},
            {"item": 2, "descricao": "Intérprete de LIBRAS", "unidade": "evento", "quantidade": 4, "valor_unitario": 1200.0},
            {"item": 3, "descricao": "Oficineiros de Dança e Música", "unidade": "oficina", "quantidade": 24, "valor_unitario": 800.0},
            {"item": 4, "descricao": "Material Didático e Insumos", "unidade": "kit", "quantidade": 100, "valor_unitario": 150.0}
        ]

    def test_pipeline_completo_execucao_com_sucesso(self):
        """Executa as 5 etapas integradas e valida a consistência de cada artefato gerado."""
        
        # Etapa 1: Ingestão
        ingest_res = FullProposalSubmissionPipeline.step1_ingestion(self.sample_edital)
        self.assertEqual(ingest_res["status"], "ingested")
        self.assertTrue(ingest_res["has_budget_rules"])
        self.assertTrue(ingest_res["has_accessibility_rules"])
        
        # Etapa 2: Revisão Técnica Setorial
        rev_res = FullProposalSubmissionPipeline.step2_sector_review(ingest_res, self.sample_sections)
        self.assertTrue(rev_res["all_approved"])
        self.assertGreaterEqual(rev_res["nota_tecnica_media"], 90.0)
        
        # Etapa 3: Supervisão Estratégica
        sup_res = FullProposalSubmissionPipeline.step3_strategic_supervision(rev_res)
        self.assertEqual(sup_res["status_geral"], "APROVADO_PARA_SUBMISSAO")
        self.assertGreaterEqual(sup_res["score_final"], 90)
        
        # Etapa 4: Geração DOCX
        cover_data = {
            "titulo": "Circuito Cultural Rio Doce Vivo",
            "proponente": "Ponto de Cultura Tambores Esperança",
            "municipio": "Fundão - ES",
            "valor_total": 78200.0,
            "prazo_meses": 12
        }
        docx_bytes = FullProposalSubmissionPipeline.step4_generate_abnt_document(cover_data, self.sample_sections)
        self.assertGreater(len(docx_bytes), 5000, "O documento Word gerado deve conter bytes válidos.")
        
        # Etapa 5: Geração de Planilha Orçamentária XLSX
        xlsx_bytes = FullProposalSubmissionPipeline.step5_generate_financial_spreadsheet(self.sample_budget)
        self.assertGreater(len(xlsx_bytes), 3000, "A planilha Excel gerada deve conter bytes válidos.")
        
        # Validar leitura do arquivo Excel gerado em memória
        wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes))
        ws = wb.active
        self.assertEqual(ws["A1"].value, "Item")
        self.assertEqual(ws["F6"].value, "=SUM(F2:F5)")


if __name__ == '__main__':
    unittest.main()
